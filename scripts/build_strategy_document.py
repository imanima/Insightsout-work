from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path("/Users/nimani/insightsout website")
OUTPUT = ROOT / "InsightsOut_Strategy_Audience_and_Partnership_Model.docx"
LOGO = Path("/Users/nimani/Downloads/insightsout-logo-gem (3).png")

# Selected preset: standard_business_brief.
# Named brand overrides: InsightsOut colors replace the preset heading colors,
# and the editorial cover uses larger display type.
FONT = "Calibri"
INK = "18202D"
MIDNIGHT = "0F141B"
BLUE = "526EF2"
PURPLE = "8761D8"
TEAL = "45BCD4"
MUTED = "5F6979"
LIGHT = "F5F7FB"
PALE_BLUE = "EEF2FF"
PALE_PURPLE = "F4F0FF"
PALE_TEAL = "ECFAFC"
WHITE = "FFFFFF"
BORDER = "D9DEE8"
GOLD = "E3A72F"

TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_font(run, name=FONT, size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def set_widow_control(paragraph, value=True):
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:widowControl"))
    if node is None:
        node = OxmlElement("w:widowControl")
        p_pr.append(node)
    node.set(qn("w:val"), "1" if value else "0")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins=CELL_MARGIN):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in margins.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    assert sum(widths_dxa) == TABLE_WIDTH_DXA
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    set_table_borders(table)


def set_cell_text(cell, text, bold=False, color=INK, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    set_widow_control(paragraph)
    run = paragraph.add_run(text)
    set_font(run, size=size, color=color, bold=bold)


def style_table_header(row, fill=MIDNIGHT):
    set_repeat_table_header(row)
    for cell in row.cells:
        shade_cell(cell, fill)
        for run in cell.paragraphs[0].runs:
            set_font(run, size=9.5, color=WHITE, bold=True)


def add_table(doc, headers, rows, widths_dxa, header_fill=MIDNIGHT, banded=True):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=WHITE, size=9.5)
    style_table_header(table.rows[0], header_fill)

    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell_width(cells[idx], widths_dxa[idx])
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_text(cells[idx], value, size=9.3)
            if banded and row_idx % 2 == 1:
                shade_cell(cells[idx], LIGHT)
        prevent_row_split(table.rows[-1])
    return table


def set_paragraph_fill_and_border(paragraph, fill, border_color=BLUE):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = borders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        borders.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border_color)


def add_callout(doc, text, fill=PALE_BLUE, border_color=BLUE, size=13, bold=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(14)
    paragraph.paragraph_format.line_spacing = 1.15
    set_paragraph_fill_and_border(paragraph, fill, border_color)
    run = paragraph.add_run(text)
    set_font(run, size=size, color=MIDNIGHT, bold=bold)
    set_widow_control(paragraph)
    return paragraph


def add_kicker(doc, text, color=PURPLE, align=WD_ALIGN_PARAGRAPH.LEFT, after=4):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    set_keep_with_next(paragraph)
    run = paragraph.add_run(text.upper())
    set_font(run, size=9.5, color=color, bold=True)
    run.font.all_caps = True
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), "16")
    run._element.get_or_add_rPr().append(spacing)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    set_keep_with_next(paragraph)
    set_widow_control(paragraph)
    return paragraph


def add_body(doc, text, bold_prefix=None, after=None, italic=False):
    paragraph = doc.add_paragraph()
    if after is not None:
        paragraph.paragraph_format.space_after = Pt(after)
    set_widow_control(paragraph)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_font(run, size=11, color=INK, bold=True)
        run = paragraph.add_run(text[len(bold_prefix):])
        set_font(run, size=11, color=INK, italic=italic)
    else:
        run = paragraph.add_run(text)
        set_font(run, size=11, color=INK, italic=italic)
    return paragraph


def add_label_value(doc, label, value, accent=BLUE):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.10
    set_widow_control(paragraph)
    label_run = paragraph.add_run(f"{label}: ")
    set_font(label_run, size=11, color=accent, bold=True)
    value_run = paragraph.add_run(value)
    set_font(value_run, size=11, color=INK)
    return paragraph


def add_numbering_definition(doc, kind):
    numbering = doc.part.numbering_part.element
    existing_abstract_ids = [
        int(x.get(qn("w:abstractNumId")))
        for x in numbering.findall(qn("w:abstractNum"))
    ]
    existing_num_ids = [
        int(x.get(qn("w:numId")))
        for x in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(existing_abstract_ids, default=-1) + 1
    num_id = max(existing_num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(lvl_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    lvl.append(suffix)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE if kind == "bullet" else INK)
    r_pr.append(color)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(doc, text, num_id, bold_prefix=None):
    paragraph = doc.add_paragraph(style="List Paragraph")
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)
    p_pr.append(num_pr)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    set_widow_control(paragraph)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_font(run, size=11, color=INK, bold=True)
        run = paragraph.add_run(text[len(bold_prefix):])
        set_font(run, size=11, color=INK)
    else:
        run = paragraph.add_run(text)
        set_font(run, size=11, color=INK)
    return paragraph


def add_page_break(doc):
    doc.add_page_break()


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_font(run, size=9, color=MUTED)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    style_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, PURPLE, 12, 6),
        "Heading 3": (12, MIDNIGHT, 8, 4),
    }
    for style_name, (size, color, before, after) in style_tokens.items():
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    list_style = doc.styles["List Paragraph"]
    list_style.font.name = FONT
    list_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    list_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    list_style.font.size = Pt(11)
    list_style.paragraph_format.left_indent = Inches(0.5)
    list_style.paragraph_format.first_line_indent = Inches(-0.25)
    list_style.paragraph_format.space_before = Pt(0)
    list_style.paragraph_format.space_after = Pt(8)
    list_style.paragraph_format.line_spacing = 1.167


def configure_sections(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Twips(708)
        section.footer_distance = Twips(708)
        section.different_first_page_header_footer = True

        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("INSIGHTSOUT  |  STRATEGY, AUDIENCE, AND PARTNERSHIP MODEL")
        set_font(run, size=8.5, color=MUTED, bold=True)

        footer = section.footer
        p = footer.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        add_page_number(p)


def add_cover(doc):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(28)
    paragraph.paragraph_format.space_after = Pt(18)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        paragraph.add_run().add_picture(str(LOGO), width=Inches(1.0))

    add_kicker(doc, "Executive strategy", color=PURPLE, align=WD_ALIGN_PARAGRAPH.CENTER, after=16)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("InsightsOut")
    set_font(run, size=20, color=MIDNIGHT, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(14)
    title.paragraph_format.line_spacing = 1.0
    run = title.add_run("Strategy, Audience,\nand Partnership Model")
    set_font(run, size=30, color=MIDNIGHT, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(38)
    run = subtitle.add_run("Human leadership through rapid change")
    set_font(run, size=15, color=BLUE, bold=False)

    descriptor = doc.add_paragraph()
    descriptor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    descriptor.paragraph_format.left_indent = Inches(0.55)
    descriptor.paragraph_format.right_indent = Inches(0.55)
    descriptor.paragraph_format.space_after = Pt(74)
    descriptor.paragraph_format.line_spacing = 1.25
    run = descriptor.add_run(
        "A single model connecting community, cohorts, private coaching, "
        "organizational AI enablement, partner delivery, and field research."
    )
    set_font(run, size=11.5, color=MUTED)

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.paragraph_format.space_after = Pt(4)
    run = date.add_run("JULY 2026")
    set_font(run, size=10, color=PURPLE, bold=True)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Prepared by InsightsOut")
    set_font(run, size=9.5, color=MUTED, italic=True)


def add_page_title(doc, kicker, title, intro=None):
    add_kicker(doc, kicker)
    add_heading(doc, title, 1)
    if intro:
        add_body(doc, intro, after=12)


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    bullet_id = add_numbering_definition(doc, "bullet")
    event_number_id = add_numbering_definition(doc, "decimal")
    lab_number_id = add_numbering_definition(doc, "decimal")

    core = doc.core_properties
    core.title = "InsightsOut Strategy, Audience, and Partnership Model"
    core.subject = "Community, programs, AI enablement, partnerships, and research"
    core.author = "InsightsOut"
    core.keywords = "InsightsOut, community, cohort, coaching, AI enablement, partnerships, research"
    core.comments = (
        "Consolidates and updates the InsightsOut business model and audience, "
        "events, and community strategy."
    )

    # Page 1: editorial cover
    add_cover(doc)
    add_page_break(doc)

    # Page 2: executive summary
    add_page_title(
        doc,
        "The model at a glance",
        "One brand. Several clear ways to move forward.",
        (
            "InsightsOut helps people and organizations respond to rapid change with "
            "more clarity, confidence, and shared ownership. We create practical spaces "
            "where people can understand what is changing, choose what matters, and act."
        ),
    )
    add_callout(
        doc,
        "The core promise: move from fear, overload, or stalled adoption to a clear next move people can own.",
        fill=PALE_BLUE,
        border_color=BLUE,
        size=13,
        bold=True,
    )
    add_heading(doc, "Five connected paths", 2)
    for label, value in [
        ("Community", "Welcoming live events that help people name the real issue and leave with a practical step."),
        ("Cohorts", "Small groups that turn reflection into action through practice, support, and accountability."),
        ("Private coaching", "Focused support for role change, leadership decisions, founder pressure, and personal direction."),
        ("Organizations and partners", "Practical AI enablement and leadership programs that connect new technology to real work."),
        ("Research", "Field notes and event based learning that improve the work and make patterns visible."),
    ]:
        add_label_value(doc, label, value)

    add_heading(doc, "Strategic choices", 2)
    choices = [
        "Build one community umbrella, not five separate communities.",
        "Lead with the human problem. Use AI as context and capability, not as the whole identity.",
        "Make every offer specific about who it is for, what happens, and what participants leave with.",
        "Treat partnerships as a core delivery channel for AI labs, learning organizations, and transformation consultancies.",
        "Use research as a learning and credibility engine, with clear methods and careful claims.",
    ]
    for item in choices:
        add_list_item(doc, item, bullet_id)
    add_page_break(doc)

    # Page 3: problem and approach
    add_page_title(
        doc,
        "Why this work matters",
        "Rapid change is not only a technology problem.",
        (
            "People are being asked to learn new tools, reconsider their roles, make faster decisions, "
            "and lead others through uncertainty. The pressure is operational, emotional, and personal."
        ),
    )
    add_heading(doc, "The problem InsightsOut solves", 2)
    for item in [
        "People feel overwhelmed by the pace of change and unsure what it means for their work or identity.",
        "Managers must communicate clearly, support anxious teams, and redesign work before every answer is known.",
        "Founders and leaders face too many tools, decisions, and expectations, which can weaken focus.",
        "Organizations can provide access and training without creating confident use, clear ownership, or lasting adoption.",
        "People often try to navigate these questions alone, even when better answers require conversation and shared learning.",
    ]:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "The InsightsOut approach", 2)
    add_body(
        doc,
        (
            "InsightsOut combines facilitated conversation, leadership development, practical experimentation, "
            "and accountability. The work begins with the real pressure people are carrying, then connects that "
            "pressure to decisions, roles, workflows, and action."
        ),
    )
    add_callout(
        doc,
        "Participation creates ownership. Practical tests create evidence. Reflection turns experience into better judgment.",
        fill=PALE_PURPLE,
        border_color=PURPLE,
        size=12.5,
    )
    add_heading(doc, "What makes the model credible", 2)
    for label, value in [
        ("Human", "The work makes room for emotion, identity, trust, and relationships without becoming vague or therapeutic."),
        ("Practical", "Every experience produces a decision, experiment, conversation, or plan that can be used immediately."),
        ("Participatory", "People help shape how change happens instead of receiving a finished answer from above."),
        ("Grounded", "Insights are tied to real events, participant feedback, completed experiments, and visible outcomes."),
    ]:
        add_label_value(doc, label, value, accent=PURPLE)
    add_page_break(doc)

    # Page 4: change journey and architecture
    add_page_title(
        doc,
        "How the work moves",
        "A repeatable journey from pressure to progress.",
        "The same human arc can support an individual, a cohort, a team, an organization, or a partner program.",
    )
    journey_rows = [
        ("1", "Name the change", "Clarify what is happening, what feels at stake, and what is still uncertain."),
        ("2", "Make sense together", "Separate evidence from assumptions and learn from other perspectives."),
        ("3", "Choose what matters", "Identify the decision, role, relationship, or workflow that needs attention now."),
        ("4", "Test a practical move", "Run a small experiment, lead a conversation, or change one part of the work."),
        ("5", "Learn and return", "Review what happened, strengthen judgment, and decide the next move."),
    ]
    add_table(doc, ["Step", "Purpose", "What happens"], journey_rows, [650, 2450, 6260], header_fill=MIDNIGHT)
    add_callout(
        doc,
        "The intended result is more agency, clearer roles, stronger leadership, and visible progress.",
        fill=PALE_TEAL,
        border_color=TEAL,
        size=12.5,
        bold=True,
    )

    add_heading(doc, "Portfolio architecture", 2)
    portfolio_rows = [
        ("Community events", "Recognize the problem", "A clear next step"),
        ("Small group cohorts", "Practice and follow through", "Evidence, support, accountability"),
        ("Private coaching", "Work through a personal decision", "Direction and a tailored plan"),
        ("Organization programs", "Connect AI to real work", "A tested use case and adoption path"),
        ("Partner programs", "Add the human adoption layer", "Better application, learning, and feedback"),
        ("The Signal", "Capture patterns from the field", "Useful research and stronger programs"),
    ]
    add_table(doc, ["Path", "Primary job", "Immediate value"], portfolio_rows, [2240, 3360, 3760], header_fill=PURPLE)
    add_page_break(doc)

    # Page 5: audiences and buyers
    add_page_title(
        doc,
        "Who the work is for",
        "Participants and buyers need different messages.",
        (
            "The person experiencing the change is not always the person paying for the work. "
            "Clear communication should name both."
        ),
    )
    add_heading(doc, "Primary participant groups", 2)
    participants = [
        ("People navigating role change", "Employed, between roles, or reconsidering their direction. They need confidence, pathfinding, and action."),
        ("Managers and people leaders", "They need to communicate through uncertainty, support emotion, clarify roles, and help teams adapt."),
        ("Founders and senior leaders", "They need focus, better decisions, and a way to lead without carrying the pressure alone."),
        ("Teams adopting AI", "They need practical learning, safe experimentation, clear boundaries, and ownership of new ways of working."),
    ]
    for label, value in participants:
        add_label_value(doc, label, value)

    add_heading(doc, "Primary buyers and sponsors", 2)
    buyers = [
        "Employers, HR, learning and development, and change leaders sponsoring team or leadership work.",
        "AI labs and product companies that want customers and teams to move from access to confident use.",
        "AI learning organizations that want application, role translation, and follow through after training.",
        "AI transformation consultancies that want a stronger human adoption workstream around strategy and implementation.",
        "Leadership and coaching organizations that want an AI era change and enablement layer.",
    ]
    for item in buyers:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "Future or partner specific extensions", 2)
    add_body(
        doc,
        (
            "Small business owners, independent professionals, students, and younger people remain valuable audiences. "
            "They should begin as adapted programs or partner sponsored pilots rather than separate public communities."
        ),
    )
    add_callout(
        doc,
        "Launch focus: role change, managers, founders and leaders, and teams putting AI into practice.",
        fill=PALE_BLUE,
        border_color=BLUE,
        size=12.5,
        bold=True,
    )
    add_page_break(doc)

    # Page 6: community and events
    add_page_title(
        doc,
        "Community and event strategy",
        "One community. Four practical reasons to join.",
        (
            "Run two public events each month and rotate the four themes. Each event should stand alone, "
            "while also leading naturally to a cohort, coaching conversation, organization program, or research activity."
        ),
    )
    event_rows = [
        (
            "My Role Is Changing. What Do I Do Next?",
            "People who feel uncertain about their job, skills, or direction",
            "A clearer view of what is changing, a map of personal value, and one pathfinding experiment",
        ),
        (
            "How Do I Lead a Team Through Constant Change?",
            "Managers and people leaders supporting anxious or unevenly adopting teams",
            "A practical team conversation, clearer roles, and one leadership action",
        ),
        (
            "Too Much Change. What Matters Now?",
            "Founders and leaders facing decision overload and constant pressure",
            "A focus map, one decision to make, and one experiment to run",
        ),
        (
            "What Should Stay Human?",
            "People making decisions about judgment, trust, care, responsibility, and delegation",
            "A boundary to protect, a principle to use, and one real decision to revisit",
        ),
    ]
    add_table(
        doc,
        ["Event", "Who recognizes themselves", "What they leave with"],
        event_rows,
        [2700, 3000, 3660],
        header_fill=MIDNIGHT,
    )

    add_heading(doc, "A consistent event experience", 2)
    for item in [
        "Arrive with a real question or pressure point.",
        "Use guided reflection to name what is happening.",
        "Learn through honest small group conversation.",
        "Apply one practical tool, map, or decision frame.",
        "Choose a take home action and a way to return with what was learned.",
    ]:
        add_list_item(doc, item, event_number_id)

    add_callout(
        doc,
        (
            "Between sessions, participants might test a workflow, learn a tool, lead a team conversation, "
            "or bring an initiative into an enablement lab. Accountability should support learning, not create pressure."
        ),
        fill=PALE_PURPLE,
        border_color=PURPLE,
        size=11.5,
    )
    add_page_break(doc)

    # Page 7: cohorts and coaching
    add_page_title(
        doc,
        "Cohorts and private coaching",
        "Two ways to turn insight into sustained movement.",
        (
            "The website should make both paths easy to understand and easy to enter. "
            "The cohort is the shared journey. Coaching is the more personal and flexible path."
        ),
    )
    add_heading(doc, "Small group cohort: Find Your Next Move", 2)
    add_body(
        doc,
        (
            "A focused group experience for people making sense of role change, leadership pressure, or a changing future of work. "
            "Participants learn from one another while completing practical work between sessions."
        ),
    )
    add_heading(doc, "What participants get", 3)
    for item in [
        "A clearer picture of what is changing and what is still within their control.",
        "A map of strengths, judgment, relationships, and value beyond a job description.",
        "A small experiment that creates evidence about a possible next direction.",
        "Support and accountability from people facing related questions.",
        "A personal plan for the next stage of work or leadership.",
    ]:
        add_list_item(doc, item, bullet_id)
    add_body(
        doc,
        "Recommended pilot: four facilitated sessions, practical work between sessions, and a final progress review.",
        italic=True,
    )

    add_heading(doc, "Private coaching with Nima Imani", 2)
    add_body(
        doc,
        (
            "Private coaching supports a specific decision or transition with more depth. "
            "It can serve people navigating role change, managers leading through uncertainty, and founders carrying complex pressure."
        ),
    )
    add_heading(doc, "What clients get", 3)
    for item in [
        "A confidential place to think clearly without performing certainty.",
        "A sharper definition of the problem beneath the noise.",
        "Direct reflection, practical tools, and decisions tied to real life.",
        "A tailored action plan and follow through.",
    ]:
        add_list_item(doc, item, bullet_id)
    add_callout(
        doc,
        (
            "Website role: keep the coaching page directly shareable and easy to book, "
            "but do not make it compete with the main InsightsOut navigation."
        ),
        fill=PALE_TEAL,
        border_color=TEAL,
        size=11.5,
    )
    add_page_break(doc)

    # Page 8: organization offer
    add_page_title(
        doc,
        "Organizational AI enablement",
        "Help people put new technology into practice.",
        (
            "Organizations do not only need awareness or tool access. They need teams to connect AI to real work, "
            "learn through use, clarify responsibility, and lead the transition well."
        ),
    )
    add_callout(
        doc,
        "The offer: a practical AI Enablement Lab supported by manager and leadership coaching.",
        fill=PALE_BLUE,
        border_color=BLUE,
        size=13,
        bold=True,
    )
    add_heading(doc, "How the lab works", 2)
    lab_steps = [
        "Choose a real workflow or business need.",
        "Define what people should own and where AI may help.",
        "Learn the relevant capability in the context of that work.",
        "Run a small, safe test with clear quality and responsibility checks.",
        "Review what worked, what did not, and what the team learned.",
        "Document the new workflow, ownership, and next adoption step.",
    ]
    for item in lab_steps:
        add_list_item(doc, item, lab_number_id)

    add_heading(doc, "What the organization receives", 2)
    outputs = [
        ("A tested use case", "Not a generic demonstration. A real workflow with evidence from use."),
        ("Clear human and AI roles", "Who decides, reviews, approves, communicates, and remains accountable."),
        ("Team agreements", "Practical expectations for quality, trust, privacy, escalation, and learning."),
        ("An adoption plan", "Next steps, owners, follow through, and what should be tested next."),
        ("Leadership support", "Tools for managers to communicate, listen, and guide role and workflow change."),
        ("Field insight", "What participants are experiencing, where adoption is blocked, and what support is needed."),
    ]
    for label, value in outputs:
        add_label_value(doc, label, value)

    add_heading(doc, "Delivery formats", 2)
    add_body(
        doc,
        "Leader workshop, team lab, multi session enablement program, manager cohort, or leadership coaching.",
    )
    add_page_break(doc)

    # Page 9: partnerships
    add_page_title(
        doc,
        "Partnerships",
        "A clear human adoption layer for strong technical and learning partners.",
        (
            "Potential partners already bring valuable products, learning, strategy, or implementation. "
            "InsightsOut complements that work by helping people apply it, lead through it, and learn from real use."
        ),
    )
    add_callout(
        doc,
        "You bring the product, learning, or transformation plan. InsightsOut helps people put it into practice.",
        fill=PALE_PURPLE,
        border_color=PURPLE,
        size=13,
        bold=True,
    )
    partner_rows = [
        (
            "AI labs and product companies",
            "Product capability, access, technical expertise",
            "Customer learning, workflow labs, manager support, adoption feedback",
        ),
        (
            "AI learning organizations",
            "Curriculum, instruction, capability building",
            "Application labs, role translation, practice, accountability",
        ),
        (
            "AI transformation consultancies",
            "Strategy, operating model, implementation, integration",
            "Human adoption workstream, leadership coaching, team participation",
        ),
        (
            "Leadership and coaching partners",
            "Leadership development, coaching, organizational learning",
            "Rapid change, AI enablement, and workflow experimentation layer",
        ),
    ]
    add_table(
        doc,
        ["Potential partner", "Partner brings", "InsightsOut adds"],
        partner_rows,
        [2300, 3000, 4060],
        header_fill=MIDNIGHT,
    )

    add_heading(doc, "Ways to work together", 2)
    for item in [
        "Co delivered customer workshops, cohorts, or enablement programs.",
        "An embedded or white label human adoption module.",
        "Workflow adoption labs paired with product or curriculum learning.",
        "Manager, champion, or community leader enablement.",
        "Qualitative adoption feedback and carefully designed field research.",
    ]:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "Strong fit indicators", 2)
    add_body(
        doc,
        (
            "A partner has a strong product, curriculum, or transformation approach and wants better application, "
            "trusted participation, manager capability, or insight into what is happening on the ground."
        ),
    )
    add_body(
        doc,
        (
            "InsightsOut is not the primary fit for a project that only requires software configuration, "
            "technical integration, or compliance training."
        ),
        italic=True,
    )
    add_page_break(doc)

    # Page 10: research
    add_page_title(
        doc,
        "Research and field learning",
        "The Signal turns lived experience into useful insight.",
        (
            "Events, cohorts, coaching themes, and organization labs reveal what people are actually experiencing. "
            "Research makes those patterns visible and improves every offer."
        ),
    )
    add_heading(doc, "What The Signal can publish", 2)
    for item in [
        "Short field notes from recurring questions and practical experiments.",
        "Event findings based on participant responses and aggregate patterns.",
        "Clear frameworks that help people make decisions or lead conversations.",
        "Partner learning briefs about adoption barriers, confidence, roles, and application.",
        "A quarterly or annual view of how people are responding to rapid change.",
    ]:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "Core questions", 2)
    questions = [
        "What is changing in people’s roles and what do they believe is at stake?",
        "Where do fear, overload, loneliness, or lack of control block movement?",
        "What helps managers communicate and create ownership?",
        "Which workflows move from curiosity to repeated use, and why?",
        "What judgment, responsibility, care, or relationships should remain clearly human?",
        "Which experiments produce useful evidence and a stronger next decision?",
    ]
    for item in questions:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "Research standards", 2)
    standards = [
        ("Consent", "Tell participants what may be collected and how it may be used."),
        ("Privacy", "Use aggregate or anonymized data unless explicit permission is given."),
        ("Method", "Explain who participated, what was asked, and what the limits are."),
        ("Claims", "Distinguish a story, a pattern, and a measured finding."),
        ("Value", "Return useful insight to participants, partners, and the wider community."),
    ]
    for label, value in standards:
        add_label_value(doc, label, value, accent=TEAL)
    add_callout(
        doc,
        "The research loop: listen, test, learn, publish, improve the next experience.",
        fill=PALE_TEAL,
        border_color=TEAL,
        size=12.5,
        bold=True,
    )
    add_page_break(doc)

    # Page 11: business model and measures
    add_page_title(
        doc,
        "Business model and proof",
        "Each offer has a clear job, buyer, and revenue path.",
        (
            "Pricing should be validated through pilots rather than invented in advance. "
            "The immediate priority is a simple offer ladder and evidence that people act."
        ),
    )
    revenue_rows = [
        ("Community events", "Individuals or sponsors", "Free, ticketed, or sponsored", "Discovery and trust"),
        ("Small group cohorts", "Participants or employers", "Cohort fee", "Practice and follow through"),
        ("Private coaching", "Individuals or employers", "Package fee", "Depth and tailored progress"),
        ("Organization programs", "Employers", "Project or retainer", "Adoption and leadership"),
        ("Partner programs", "Labs, learning organizations, consultancies", "Project, retainer, embedded, or shared program", "Distribution and co delivery"),
        ("Research", "Community, partners, or sponsors", "Primarily learning and credibility; selective sponsored work", "Insight and stronger programs"),
    ]
    add_table(
        doc,
        ["Offer", "Primary buyer", "Revenue path", "Strategic job"],
        revenue_rows,
        [1900, 2100, 2860, 2500],
        header_fill=PURPLE,
    )

    add_heading(doc, "Measures that matter", 2)
    metrics = [
        ("Reach", "Registrations, attendance, participant mix, and return participation."),
        ("Immediate value", "People leave with a clear action, conversation, experiment, or decision."),
        ("Follow through", "Participants complete the work and return with learning."),
        ("Conversion", "Events lead to cohort applications, fit calls, coaching, or organization conversations."),
        ("Adoption", "Teams test a real workflow, repeat useful behavior, and clarify ownership."),
        ("Leadership", "Managers report stronger confidence and lead better team conversations."),
        ("Proof", "Specific testimonials, case studies, partner renewals, and documented outcomes."),
    ]
    for label, value in metrics:
        add_label_value(doc, label, value)
    add_page_break(doc)

    # Page 12: priorities
    add_page_title(
        doc,
        "Recommended next ninety days",
        "Focus the brand, test the model, and build proof.",
        (
            "The goal is not to launch everything at once. The goal is to make the core paths real, "
            "learn quickly, and show potential participants and partners what InsightsOut can do."
        ),
    )
    priority_rows = [
        (
            "1. Clarify",
            "Finalize the core website message, event pages, cohort promise, organization offer, partner page, and booking flow.",
            "Every visitor can identify the right path in under a minute.",
        ),
        (
            "2. Test",
            "Run two public events each month, pilot one cohort, and deliver one organizational enablement lab.",
            "Real participant feedback, completed experiments, and early proof.",
        ),
        (
            "3. Partner",
            "Hold focused conversations with AI labs, learning organizations, transformation consultancies, and leadership partners.",
            "A short list of strong fit partners and one co delivery pilot.",
        ),
        (
            "4. Learn",
            "Use a simple consent based feedback system and publish the first Signal field note.",
            "A visible learning loop and a more credible public voice.",
        ),
        (
            "5. Refine",
            "Set pricing, package the offers, collect testimonials, and document one strong case study.",
            "A repeatable model ready for broader outreach.",
        ),
    ]
    add_table(
        doc,
        ["Priority", "What to do", "What success looks like"],
        priority_rows,
        [1500, 4660, 3200],
        header_fill=MIDNIGHT,
    )

    add_heading(doc, "Decisions still to make", 2)
    for item in [
        "Final pricing and package length for cohorts, coaching, organization work, and partner delivery.",
        "The first cohort dates, audience, and application process.",
        "Which technical or training capabilities InsightsOut will deliver directly and which will remain partner led.",
        "The consent, privacy, and research process for events and programs.",
        "The proof standard for testimonials, case studies, and public outcome claims.",
    ]:
        add_list_item(doc, item, bullet_id)

    add_callout(
        doc,
        (
            "People do not need more noise about the future. They need a place to understand what is changing, "
            "people to think with, and a practical next move. InsightsOut can be that place."
        ),
        fill=PALE_BLUE,
        border_color=BLUE,
        size=13,
        bold=True,
    )
    add_body(
        doc,
        (
            "Document note: this strategy consolidates and updates the InsightsOut business model and audience, "
            "events, and community working documents. Partner examples describe target relationship types, not confirmed partnerships."
        ),
        italic=True,
        after=0,
    )

    doc.save(OUTPUT)
    return doc


def audit_document(doc):
    errors = []
    section = doc.sections[0]
    expected = {
        "page_width": Inches(8.5),
        "page_height": Inches(11),
        "top_margin": Inches(1),
        "bottom_margin": Inches(1),
        "left_margin": Inches(1),
        "right_margin": Inches(1),
        "header_distance": Twips(708),
        "footer_distance": Twips(708),
    }
    for name, value in expected.items():
        if abs(getattr(section, name) - value) > 5:
            errors.append(f"section {name} does not match preset")

    normal = doc.styles["Normal"]
    if normal.font.name != FONT or normal.font.size != Pt(11):
        errors.append("Normal style font or size does not match preset")
    if normal.paragraph_format.space_after != Pt(6):
        errors.append("Normal style after spacing does not match preset")

    expected_styles = {
        "Heading 1": (Pt(16), Pt(16), Pt(8)),
        "Heading 2": (Pt(13), Pt(12), Pt(6)),
        "Heading 3": (Pt(12), Pt(8), Pt(4)),
    }
    for style_name, values in expected_styles.items():
        style = doc.styles[style_name]
        if style.font.size != values[0]:
            errors.append(f"{style_name} size does not match preset")
        if style.paragraph_format.space_before != values[1]:
            errors.append(f"{style_name} before spacing does not match preset")
        if style.paragraph_format.space_after != values[2]:
            errors.append(f"{style_name} after spacing does not match preset")

    for index, table in enumerate(doc.tables, start=1):
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        layout = tbl_pr.find(qn("w:tblLayout"))
        if tbl_w is None or tbl_w.get(qn("w:w")) != str(TABLE_WIDTH_DXA):
            errors.append(f"table {index} width is not {TABLE_WIDTH_DXA} DXA")
        if tbl_ind is None or tbl_ind.get(qn("w:w")) != str(TABLE_INDENT_DXA):
            errors.append(f"table {index} indent is not {TABLE_INDENT_DXA} DXA")
        if layout is None or layout.get(qn("w:type")) != "fixed":
            errors.append(f"table {index} layout is not fixed")
        grid_widths = [
            int(col.get(qn("w:w")))
            for col in table._tbl.tblGrid.findall(qn("w:gridCol"))
        ]
        if sum(grid_widths) != TABLE_WIDTH_DXA:
            errors.append(f"table {index} grid does not total {TABLE_WIDTH_DXA} DXA")
        for row in table.rows:
            if len(row.cells) != len(grid_widths):
                errors.append(f"table {index} has merged or irregular row geometry")
                break
            for cell_index, cell in enumerate(row.cells):
                tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                if tc_w is None or int(tc_w.get(qn("w:w"))) != grid_widths[cell_index]:
                    errors.append(f"table {index} has a cell width that does not match tblGrid")
                    break

    with ZipFile(OUTPUT) as archive:
        numbering_xml = archive.read("word/numbering.xml").decode("utf-8")
        if 'w:left="720"' not in numbering_xml or 'w:hanging="360"' not in numbering_xml:
            errors.append("numbering definitions do not carry preset list indents")
        if 'w:after="160"' not in numbering_xml or 'w:line="280"' not in numbering_xml:
            errors.append("numbering definitions do not carry preset list spacing")

    return errors


if __name__ == "__main__":
    document = build_document()
    audit_errors = audit_document(document)
    print(f"Created: {OUTPUT}")
    print(f"Paragraphs: {len(document.paragraphs)}")
    print(f"Tables: {len(document.tables)}")
    print(f"Sections: {len(document.sections)}")
    if audit_errors:
        print("AUDIT ERRORS")
        for error in audit_errors:
            print(f"  {error}")
        raise SystemExit(1)
    print("Preset and geometry audit: passed")
