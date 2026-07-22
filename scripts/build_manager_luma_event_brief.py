from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/nimani/insightsout website")
OUTPUT = ROOT / "InsightsOut_Manager_Workshop_Luma_Event_Brief.docx"
LOGO = Path("/Users/nimani/Downloads/insightsout-logo-gem (3).png")

# Selected preset: launch_messaging_guide, based on compact_reference_guide.
# Selected first page pattern: workshop_agenda.
# Named brand overrides: InsightsOut colors and gem logo.
FONT = "Calibri"
INK = "17202D"
MIDNIGHT = "0F141B"
BLUE = "526EF2"
PURPLE = "8761D8"
TEAL = "45BCD4"
MUTED = "606A79"
LIGHT = "F5F7FB"
PALE_BLUE = "EEF2FF"
PALE_PURPLE = "F4F0FF"
PALE_TEAL = "ECFAFC"
WHITE = "FFFFFF"
BORDER = "D9DEE8"
GOLD = "E3A72F"

TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN = {"top": 100, "bottom": 100, "start": 140, "end": 140}


def rgb(value):
    return RGBColor.from_string(value)


def set_font(run, name=FONT, size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def set_keep_together(paragraph, value=True):
    paragraph.paragraph_format.keep_together = value


def set_widow_control(paragraph, value=True):
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:widowControl"))
    if node is None:
        node = OxmlElement("w:widowControl")
        p_pr.append(node)
    node.set(qn("w:val"), "1" if value else "0")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins=CELL_MARGIN):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in margins.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size=6, inside=True):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    edges = ["top", "left", "bottom", "right"]
    if inside:
        edges += ["insideH", "insideV"]
    for edge in edges:
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA, borders=True):
    assert sum(widths_dxa) == TABLE_WIDTH_DXA
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if borders:
        set_table_borders(table)


def set_cell_text(cell, text, bold=False, color=INK, size=9.6, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    set_widow_control(p)
    run = p.add_run(text)
    set_font(run, size=size, color=color, bold=bold)


def add_table(doc, headers, rows, widths_dxa, header_fill=MIDNIGHT, banded=True):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, color=WHITE, size=9.4)
        shade_cell(table.rows[0].cells[index], header_fill)
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell_width(cells[index], widths_dxa[index])
            set_cell_margins(cells[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_text(cells[index], value, size=9.3)
            if banded and row_index % 2 == 1:
                shade_cell(cells[index], LIGHT)
        prevent_row_split(table.rows[-1])
    add_spacer(doc, 6)
    return table


def add_metric_strip(doc, items):
    widths = [TABLE_WIDTH_DXA // len(items)] * len(items)
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    table = doc.add_table(rows=1, cols=len(items))
    set_table_geometry(table, widths, borders=False)
    set_table_borders(table, color="D9DFFB", size=6, inside=True)
    set_repeat_table_header(table.rows[0])
    for index, (value, label) in enumerate(items):
        cell = table.rows[0].cells[index]
        shade_cell(cell, PALE_BLUE if index != 1 else PALE_PURPLE)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(value)
        set_font(r, size=14, color=MIDNIGHT, bold=True)
        r = p.add_run("\n" + label.upper())
        set_font(r, size=8.2, color=MUTED, bold=True)
    add_spacer(doc, 12)


def add_numbering_definition(doc, kind):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def add_bullet(doc, text, num_id, bold_prefix=None):
    p = doc.add_paragraph()
    apply_numbering(p, num_id)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_widow_control(p)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, size=11, color=INK, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_font(r, size=11, color=INK)
    else:
        r = p.add_run(text)
        set_font(r, size=11, color=INK)
    return p


def add_numbered(doc, text, num_id, bold_prefix=None):
    return add_bullet(doc, text, num_id, bold_prefix)


def add_spacer(doc, points):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(points)
    p.paragraph_format.line_spacing = 0.2


PAGE_BREAK_PENDING = False


def add_kicker(doc, text, color=PURPLE, after=5):
    global PAGE_BREAK_PENDING
    p = doc.add_paragraph()
    if PAGE_BREAK_PENDING:
        p.paragraph_format.page_break_before = True
        PAGE_BREAK_PENDING = False
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    set_keep_with_next(p)
    r = p.add_run(text.upper())
    set_font(r, size=9.2, color=color, bold=True)
    r.font.all_caps = True
    return p


def add_section_title(doc, kicker, title, new_page=True):
    p = doc.add_paragraph()
    if new_page:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(24 if new_page else 0)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.05
    set_keep_with_next(p)
    set_keep_together(p)
    r = p.add_run(kicker.upper())
    set_font(r, size=9.2, color=TEAL, bold=True)
    r.add_break()
    r = p.add_run(title)
    set_font(r, size=16, color=PURPLE, bold=True)
    return p


def add_cover_brand(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    if LOGO.exists():
        r = p.add_run()
        picture = r.add_picture(str(LOGO), width=Inches(0.24))
        picture._inline.docPr.set("title", "InsightsOut logo")
        picture._inline.docPr.set("descr", "InsightsOut multicolor gem logo")
        r = p.add_run("  InsightsOut")
        set_font(r, size=11, color=MIDNIGHT, bold=True)
    return p


def add_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 0.98
    set_keep_with_next(p)
    r = p.add_run(text)
    set_font(r, size=30, color=MIDNIGHT, bold=True)
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.line_spacing = 1.18
    set_keep_with_next(p)
    r = p.add_run(text)
    set_font(r, size=13.5, color=MUTED)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_keep_with_next(p)
    set_widow_control(p)
    return p


def add_body(doc, text, bold_prefix=None, italic=False, after=None):
    p = doc.add_paragraph()
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    set_widow_control(p)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, size=11, color=INK, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_font(r, size=11, color=INK, italic=italic)
    else:
        r = p.add_run(text)
        set_font(r, size=11, color=INK, italic=italic)
    return p


def set_paragraph_fill_and_border(paragraph, fill, border_color=BLUE):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = borders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        borders.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border_color)


def add_callout(doc, text, label=None, fill=PALE_BLUE, border_color=BLUE, size=12.2):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(13)
    p.paragraph_format.line_spacing = 1.18
    set_paragraph_fill_and_border(p, fill, border_color)
    set_keep_together(p)
    if label:
        r = p.add_run(label.upper() + "\n")
        set_font(r, size=8.8, color=border_color, bold=True)
    r = p.add_run(text)
    set_font(r, size=size, color=MIDNIGHT, bold=False)
    return p


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.24)
    p.paragraph_format.right_indent = Inches(0.24)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(14)
    p.paragraph_format.line_spacing = 1.18
    set_paragraph_fill_and_border(p, PALE_PURPLE, PURPLE)
    r = p.add_run("“" + text + "”")
    set_font(r, size=13, color=MIDNIGHT, italic=True)
    return p


def add_label_value(doc, label, value, accent=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.15
    set_widow_control(p)
    r = p.add_run(label + ": ")
    set_font(r, size=11, color=accent, bold=True)
    r = p.add_run(value)
    set_font(r, size=11, color=INK)
    return p


def add_page_break(doc):
    # Content is allowed to flow naturally. This avoids sparse pages when a
    # preceding section already ends at the bottom of a page.
    return None


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Heading 1": (16, PURPLE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, "2F527A", 10, 5),
    }
    for name, (size, color, before, after) in tokens.items():
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr)
    run._r.append(fld_end)
    set_font(run, size=8.5, color=MUTED)


def set_running_header_footer(section):
    for header in (section.header, section.even_page_header):
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.clear()
        p.paragraph_format.space_after = Pt(0)

    for footer in (section.footer, section.even_page_footer):
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run("InsightsOut   •   Human leadership through rapid change   •   ")
        set_font(r, size=8.3, color=MUTED)
        add_page_field(p)


def configure_page(doc):
    section = doc.sections[0]
    doc.settings.odd_and_even_pages_header_footer = True
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)
    section.different_first_page_header_footer = False

    set_running_header_footer(section)


doc = Document()
configure_page(doc)
configure_styles(doc)
bullet_id = add_numbering_definition(doc, "bullet")
message_number_id = add_numbering_definition(doc, "decimal")
registration_number_id = add_numbering_definition(doc, "decimal")
funnel_number_id = add_numbering_definition(doc, "decimal")

# Cover and recommendation
add_cover_brand(doc)
add_kicker(doc, "Manager Workshop Launch Brief", TEAL)
add_title(doc, "What Do I Say to My Team Right Now?")
add_subtitle(doc, "A Luma ready event plan for managers leading through pressure, mixed reactions, rapid change, and AI adoption.")
add_metric_strip(doc, [("90 min", "interactive session"), ("12 to 24", "managers"), ("1 real", "team situation")])
add_callout(
    doc,
    "Managers do not need another presentation about change. They need a place to understand what their team is reacting to, decide what they can influence, and prepare one honest conversation.",
    "Core promise",
    PALE_TEAL,
    TEAL,
    12.4,
)
add_heading(doc, "Recommended public positioning", 1)
add_body(doc, "Lead with the manager’s immediate question, not the technology. The title names the conversation they are worried about. The subtitle and event description make clear that AI adoption and changing ways of working are part of the pressure.")
add_label_value(doc, "Recommended title", "What Do I Say to My Team Right Now?")
add_label_value(doc, "Public subtitle", "A practical workshop for managers leading teams through uncertainty, pressure, and changing ways of working.")
add_label_value(doc, "Internal descriptor", "A manager leadership lab for team resistance, communication, and responsible AI adoption.")
add_label_value(doc, "Primary next step", "A manager cohort, an organization pilot, or a private conversation about the team’s adoption challenge.")
add_callout(doc, "The event should feel useful even if the participant never buys anything. That credibility is what makes the next conversation natural.", "Design principle", PALE_PURPLE, PURPLE, 11.4)

# Audience and pain
add_section_title(doc, "01  Audience and need", "The manager in the middle")
add_body(doc, "This workshop is for managers, team leads, and heads of function who are expected to keep work moving while direction, tools, roles, and expectations continue to change.")
add_quote(doc, "I am expected to create clarity for my team, but I do not have all the answers myself.")
add_heading(doc, "Who this is for", 2)
for item in [
    "First line and mid level managers translating senior decisions into daily work.",
    "Team leads introducing new tools, workflows, or expectations while continuing to deliver.",
    "People managers hearing anxiety, skepticism, frustration, silence, or quiet avoidance from their teams.",
    "Managers who want adoption to move forward without dismissing valid concerns or forcing enthusiasm.",
    "Leaders of technical and nontechnical teams. No AI expertise is required.",
]:
    add_bullet(doc, item, bullet_id)
add_heading(doc, "What they are carrying", 2)
for item in [
    "Communication pressure. They must explain a change before the message feels settled.",
    "Change fatigue. The team hears one more initiative while the existing workload has not gone away.",
    "Mixed reactions. Some people are curious, some afraid, some skeptical, and some appear to comply while avoiding the change.",
    "Adoption pressure. Leadership wants visible progress, but use remains uneven and the value is not always clear.",
    "Role uncertainty. People are asking what will change, what they still own, and whether their experience will continue to matter.",
    "Personal uncertainty. The manager is learning too, but feels expected to perform confidence.",
    "Human accountability. Even when a tool produces the work, the manager remains responsible for quality, relationships, judgment, and consequences.",
]:
    add_bullet(doc, item, bullet_id, item.split(". ")[0] + ". ")
add_callout(doc, "The real problem is not simply that people resist change. The manager often lacks a practical way to read the reaction, respond accurately, and create a next step the team can understand and shape.", "Problem statement", PALE_BLUE, BLUE, 11.7)

# Resistance signals
add_section_title(doc, "02  The problem beneath the problem", "Resistance is information")
add_body(doc, "Public messaging can use familiar language such as pushback or resistance. Inside the workshop, the facilitator should help participants become more precise. One behavior can point to several different needs.")
add_table(
    doc,
    ["Signal", "What may sit underneath", "A more useful manager response"],
    [
        ["Clarity", "The reason, expectation, timing, or boundary is unclear.", "Name what is decided, what is moving, and what happens next."],
        ["Capacity", "New work has been added without anything stopping.", "Remove, sequence, or narrow the change before asking for more."],
        ["Capability", "People lack practice, context, or safe support.", "Learn inside a real workflow and normalize early mistakes."],
        ["Trust", "People doubt the intent, tool, data use, or fairness.", "Invite specifics and make safeguards and tradeoffs visible."],
        ["Agency", "The change is being done to people rather than shaped with them.", "Clarify what is fixed and what the team can genuinely influence."],
        ["Identity", "Experience, craft, status, or livelihood feels threatened.", "Name the loss and show where human value and responsibility remain."],
        ["Risk", "There is a real concern about quality, privacy, bias, or harm.", "Treat the concern as useful risk input. Define review and stop conditions."],
    ],
    [1500, 3450, 4410],
    header_fill=MIDNIGHT,
)
add_heading(doc, "What managers often try first", 2)
for item in [
    "Repeat the leadership message more forcefully.",
    "Explain the benefits before understanding the concern.",
    "Treat hesitation as a mindset problem.",
    "Ask for participation when the real decision is already closed.",
    "Promise certainty, control, or job safety they cannot guarantee.",
    "Measure activity while missing trust, quality, workload, and workarounds.",
]:
    add_bullet(doc, item, bullet_id)
add_callout(doc, "The workshop replaces persuasion with better diagnosis, clearer decisions, honest communication, and a small responsible experiment.", "Shift", PALE_TEAL, TEAL, 11.7)

# Titles
add_section_title(doc, "03  Title and message direction", "Title options")
add_body(doc, "The best title should feel immediately familiar to a manager, avoid blaming the team, and remain relevant beyond a single tool rollout.")
add_table(
    doc,
    ["Title", "Why it works", "Watch for"],
    [
        ["What Do I Say to My Team Right Now?", "Immediate, human, and specific. Centers the conversation the manager is already carrying.", "Needs a subtitle to signal change and AI adoption."],
        ["My Team Is Overwhelmed. How Do I Lead?", "Strong emotional recognition. Fits managers facing change fatigue and heavy workloads.", "Less direct about resistance and communication."],
        ["Leading Change When You Do Not Have All the Answers", "Honest and credible. Removes the pressure to perform certainty.", "More familiar and less distinctive."],
        ["When Your Team Is Not Ready for More Change", "Clearly names change fatigue and the manager’s dilemma.", "Can sound as if the team is the problem."],
        ["How to Lead Through Pushback Without Forcing Buy In", "Directly addresses resistance and the fear of becoming coercive.", "Longer and more tactical in tone."],
    ],
    [3350, 3590, 2420],
    header_fill=PURPLE,
)
add_callout(doc, "Use What Do I Say to My Team Right Now? for the first public pilot. It names the pain in plain language and leaves room for participants facing AI adoption, restructuring, workflow change, or broader uncertainty.", "Recommendation", PALE_PURPLE, PURPLE, 11.7)
add_heading(doc, "Message hierarchy", 2)
for item in [
    "Title: the urgent personal question.",
    "Subtitle: who the workshop is for and the situation they are facing.",
    "First paragraph: the pressure of leading while answers are still moving.",
    "Second paragraph: team reactions, communication, and AI adoption.",
    "Promise: one clearer position, one prepared conversation, and one next move.",
]:
    add_numbered(doc, item, message_number_id)
add_heading(doc, "Language to use", 2)
add_body(doc, "Use: pressure, mixed reactions, honest conversation, changing work, what is known, what is moving, what people can shape, human responsibility, next move.")
add_body(doc, "Avoid: overcome resistance, drive compliance, future proof, transformation journey, unlock potential, human capital, stakeholder alignment, change champions.")

# Luma listing
add_section_title(doc, "04  Luma listing", "Ready to paste")
add_label_value(doc, "Event title", "What Do I Say to My Team Right Now?")
add_label_value(doc, "Subtitle", "A practical workshop for managers leading teams through uncertainty, pressure, and changing ways of working.")
add_heading(doc, "Short preview", 2)
add_body(doc, "Your team has questions. You may not have every answer. Bring one real change you are leading and leave with a clearer message, a better question, and a practical next step.")
add_heading(doc, "Full event description", 2)
add_body(doc, "Managers are often asked to create clarity while the direction, tools, roles, and expectations are still changing.")
add_body(doc, "Your team may be anxious, skeptical, tired of change, quietly avoiding a new way of working, or asking questions you cannot fully answer. At the same time, leadership expects progress. This is especially visible when organizations introduce AI into everyday work, but the challenge is larger than the technology. It is about trust, communication, workload, identity, and responsibility.")
add_body(doc, "This interactive workshop gives managers a small, honest place to work through one live situation. We will look beneath the label of resistance, identify what the team’s reaction may be telling you, sort the decisions you are carrying, and prepare one conversation you can actually have.")
add_body(doc, "You do not need to perform certainty or arrive with a perfect plan. You only need to bring something real.")
add_heading(doc, "You will leave with", 2)
for item in [
    "A clearer view of what is known, what is still moving, and what your team may be experiencing.",
    "A more useful interpretation of one team reaction and a question you can ask next.",
    "A decision you can make, involve others in, test, or escalate.",
    "A prepared opening for one honest team conversation.",
    "One leadership move to complete within 30 days.",
]:
    add_bullet(doc, item, bullet_id)

add_section_title(doc, "04  Luma listing continued", "Who should attend")
for item in [
    "Managers and team leads guiding people through changing work.",
    "Heads of function translating an AI or change initiative into daily practice.",
    "People leaders facing skepticism, anxiety, silence, uneven use, or change fatigue.",
    "Managers who want to move adoption forward without dismissing valid concerns.",
]:
    add_bullet(doc, item, bullet_id)
add_heading(doc, "What this is not", 2)
for item in [
    "A product demonstration or prompt writing class.",
    "A lecture about making people more positive.",
    "A script for forcing agreement or enthusiasm.",
    "A space to discuss change in the abstract. Bring one real situation.",
]:
    add_bullet(doc, item, bullet_id)
add_heading(doc, "Format line", 2)
add_callout(doc, "90 minutes  •  Small group  •  Guided reflection  •  Peer conversation  •  Practical rehearsal  •  Stories stay, learning leaves  •  Sharing is always optional", None, PALE_BLUE, BLUE, 11.2)
add_heading(doc, "Host bio", 2)
add_body(doc, "Nima Imani is the founder of InsightsOut, an ICF ACC coach, and a trained facilitator with more than a decade of experience across software engineering, data, AI, solution architecture, consulting, and technology leadership, including EY and Neo4j. He helps managers and teams think clearly, communicate honestly, and act responsibly when work changes quickly.")
add_heading(doc, "Recommended pilot settings", 2)
add_table(
    doc,
    ["Element", "Recommendation"],
    [
        ["Length", "90 minutes for the public pilot. Protect the practice and commitment at the end."],
        ["Group size", "12 to 24 participants. Cap registration because the small group is part of the value."],
        ["Price", "$25 founding pilot. Use free admission when a partner sponsors the room or audience."],
        ["Mode", "Choose one mode per event. In person in San Francisco is strongest for trust. Live online expands reach."],
    ],
    [1900, 7460],
    header_fill=BLUE,
)

# Experience design
add_section_title(doc, "05  Participant experience", "A 90 minute working session")
add_body(doc, "The session moves from pressure to sense making to action. It stays close to one real team situation and protects enough time for participants to practice what they will say.")
add_table(
    doc,
    ["Time", "Segment", "What happens", "Visible output"],
    [
        ["0 to 10", "Arrive and contract", "Set confidentiality and participation agreements. Choose one live situation.", "One situation selected"],
        ["10 to 25", "Name the pressure", "Private reflection and paired listening. Separate the situation from the pressure around it.", "The real leadership tension"],
        ["25 to 45", "Read the reaction", "Describe what the team is doing without assigning motive. Explore what the reaction may signal.", "One better question"],
        ["45 to 60", "Choose the lane", "Sort the decision into decide, involve, test, or escalate.", "A decision and stance"],
        ["60 to 82", "Prepare the conversation", "Write and rehearse a clear opening with one difficult question from a peer.", "A practiced team message"],
        ["82 to 90", "Commit and close", "Choose one action within 72 hours and a review point within 30 days.", "Owner, action, review date"],
    ],
    [1050, 1800, 4210, 2300],
    header_fill=MIDNIGHT,
)
add_heading(doc, "The honest team message", 2)
for item in [
    "Why this is happening. Connect the change to the work, not to hype or inevitability.",
    "What we know. State the facts, decisions, and current expectation.",
    "What we do not know. Name what is moving and when more may be known.",
    "What changes now. Be specific about behavior, workflow, timing, and what does not change.",
    "What stays human. Name judgment, care, review, escalation, and accountability.",
    "What the team can shape. Ask a real question. Do not offer influence where none exists.",
    "What happens next. Set the experiment, owner, feedback channel, and review date.",
]:
    add_bullet(doc, item, bullet_id, item.split(". ")[0] + ". ")
add_callout(doc, "Success means every participant can name what they will say, what they are inviting the team to shape, what remains human owned, and when they will review the next step.", "Session success test", PALE_TEAL, TEAL, 11.5)

# Registration and communication
add_section_title(doc, "06  Registration and participant communication", "Luma registration questions")
add_body(doc, "Keep registration light. Ask only what improves the room or the follow up.")
for item in [
    "What is your role and what kind of team do you lead?",
    "What change, new tool, or new way of working is affecting your team right now?",
    "What reaction or conversation are you finding difficult?",
    "What would make this session useful for you?",
    "Do you have any access needs or participation preferences we should know about?",
]:
    add_numbered(doc, item, registration_number_id)
add_heading(doc, "Confirmation message", 2)
add_callout(doc, "You are registered. Before the workshop, choose one real situation involving a team reaction, a change you are communicating, or a decision you are carrying. You will not be required to share anything you want to keep private. Bring a notebook and arrive ready to work with the situation, not perform an answer.", None, PALE_BLUE, BLUE, 10.8)
add_heading(doc, "Reminder message", 2)
add_callout(doc, "Tomorrow we will work with one real situation from your team. Bring the facts you know, the parts that are still moving, and one conversation you have been avoiding or rehearsing in your head. This is a participatory session, not a webinar. Please join from a place where you can reflect and speak privately.", None, PALE_PURPLE, PURPLE, 10.8)
add_heading(doc, "Frequently asked questions", 2)
add_label_value(doc, "Do I need AI expertise", "No. This is a leadership and communication workshop. AI adoption is one important context, but the work begins with your team situation.")
add_label_value(doc, "Do I need to share confidential details", "No. You can change names and identifying details, and sharing is always optional.")
add_label_value(doc, "Is this a tool training", "No. We will not teach prompts or demonstrate products.")
add_label_value(doc, "Can my whole management team attend", "Yes. For larger internal groups, use the organization version with sponsor discovery and a follow up review.")

# Promotion
add_section_title(doc, "07  Promotion", "LinkedIn launch post")
add_callout(doc, "Managers are often expected to create clarity before they have it themselves.\n\nA team may be anxious, skeptical, tired of change, or quietly avoiding a new way of working. Leadership still expects progress. The manager is left in the middle, trying to communicate honestly without making promises they cannot keep.\n\nI am hosting a small working session for managers carrying that pressure.\n\nWhat Do I Say to My Team Right Now?\n\nBring one real situation. We will look at what the team’s reaction may be telling you, sort the decision you are carrying, and prepare one conversation you can actually have.\n\nThis is not a product demo or a lecture about overcoming resistance. It is a practical place to think, practice, and choose a next move.\n\nSmall group. 90 minutes. Details and registration on Luma.", None, PALE_BLUE, BLUE, 10.6)
add_heading(doc, "Short invitation", 2)
add_body(doc, "Your team has questions. You may not have every answer. Join a small practical workshop for managers leading through pressure, mixed reactions, and changing ways of working. Bring one real situation and leave ready for one honest conversation.")
add_heading(doc, "Email invitation subject options", 2)
for item in [
    "What do I say to my team right now?",
    "A small workshop for managers carrying change",
    "When your team has questions you cannot fully answer",
]:
    add_bullet(doc, item, bullet_id)
add_heading(doc, "Partner invitation line", 2)
add_body(doc, "InsightsOut helps managers translate AI strategy and learning into daily team practice by strengthening the communication, judgment, trust, and follow through that adoption requires.")
add_heading(doc, "Visual direction for Luma", 2)
add_body(doc, "Use a real photograph of people in conversation, not an abstract technology image. Keep the graphic quiet and premium. Show the InsightsOut gem, the title, and a short line such as A practical workshop for managers. Avoid robots, glowing brains, dashboards, and dense text.")

# Funnel and research
add_section_title(doc, "08  Follow through, conversion, and learning", "The event is the beginning")
add_body(doc, "The workshop should create immediate value and reveal the next level of support without turning the room into a sales presentation.")
for item in [
    "Workshop. The manager understands one reaction, prepares one conversation, and chooses one action.",
    "Manager circle. A small cohort returns to compare what happened, practice new situations, and build accountability.",
    "Organization pilot. InsightsOut supports a manager group or team through one real adoption challenge over several weeks.",
    "Private conversation. The manager or sponsor discusses the pressure, team context, and clearest place to begin.",
]:
    add_numbered(doc, item, funnel_number_id, item.split(". ")[0] + ". ")
add_heading(doc, "Closing invitation in the room", 2)
add_callout(doc, "If this conversation helped you see the situation more clearly, there are two ways to continue. You can join a small manager circle to keep practicing with peers, or bring InsightsOut into your organization to support a manager group or team through a live change. I will share both options in the follow up. There is nothing you need to decide today.", None, PALE_PURPLE, PURPLE, 10.9)
add_heading(doc, "Follow up email", 2)
add_callout(doc, "Thank you for bringing a real situation into the room. Within the next 72 hours, try the conversation or action you prepared. Then notice what became clearer, what remained difficult, and what your team showed you.\n\nIf you want continued peer practice, reply with MANAGER CIRCLE. If your organization is working through an adoption challenge and wants support for managers or teams, reply with ORGANIZATION. You can also book a private conversation with Nima.\n\nStories stay. The shared learning can continue.", None, PALE_BLUE, BLUE, 10.6)
add_heading(doc, "Learning to capture", 2)
for item in [
    "What kinds of change are managers leading?",
    "Which reactions are most common and what appears to sit underneath them?",
    "Where do managers lack decision authority or clear information?",
    "Which conversations are managers postponing?",
    "What action do participants take within 72 hours?",
    "What changed by the 30 day review?",
]:
    add_bullet(doc, item, bullet_id)
add_body(doc, "Use only aggregated, deidentified themes in research or public writing unless a participant gives explicit permission for their story or words to be used.")

add_heading(doc, "Measures that matter", 2)
add_table(
    doc,
    ["Stage", "Useful evidence"],
    [
        ["Demand", "Registrations by role, source, and problem. Attendance rate. Waitlist demand."],
        ["Experience", "Participants can name one better question, one decision, and one conversation."],
        ["Action", "The planned conversation or experiment happens within 72 hours."],
        ["Continued need", "Replies for the manager circle, organization inquiry, or private conversation."],
        ["Research", "Repeated themes are captured without exposing personal or company details."],
    ],
    [1900, 7460],
    header_fill=TEAL,
)

# Final launch checklist
add_section_title(doc, "09  Launch checklist", "Ready for the first pilot", new_page=False)
add_heading(doc, "Before publishing", 2)
for item in [
    "Choose the date, mode, venue, capacity, and pilot price.",
    "Create a simple Luma cover using a real community or workshop photograph.",
    "Paste the recommended title, subtitle, description, host bio, and format line.",
    "Add the five registration questions and a clear consent choice for updates.",
    "Set the confirmation and reminder messages.",
    "Prepare the participant worksheet and the facilitator timing sheet.",
]:
    add_bullet(doc, item, bullet_id)
add_heading(doc, "After publishing", 2)
for item in [
    "Share the LinkedIn launch post and send the short invitation to relevant managers.",
    "Invite two or three trusted partners to share the event with manager communities.",
    "Review registrations for the language people use about their situation.",
    "Adapt examples to the group without changing the core promise.",
    "Send the follow up within 24 hours and check for action after 72 hours.",
]:
    add_bullet(doc, item, bullet_id)
add_callout(doc, "What Do I Say to My Team Right Now? is the recommended first title. It is human, urgent, and broad enough to attract managers facing AI adoption without making technology the center of the invitation.", "Final recommendation", PALE_TEAL, TEAL, 12)

doc.core_properties.title = "InsightsOut Manager Workshop Luma Event Brief"
doc.core_properties.subject = "Luma event copy, manager pain points, workshop design, promotion, follow through, and research"
doc.core_properties.author = "InsightsOut"
doc.core_properties.keywords = "InsightsOut, managers, change, AI adoption, team resistance, workshop, Luma"
doc.save(OUTPUT)
print(OUTPUT)
